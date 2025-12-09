# python forensics_reporter.py --pdf sample.pdf --images images_dir --out report.docx

import argparse
import io
import csv
import hashlib
import zipfile
import tempfile
import datetime
import os
import sys
import logging
from pathlib import Path
from typing import Dict, Optional, Set

from docx import Document
from docx.shared import Inches, Pt
import fitz  # PyMuPDF
from PIL import Image, ExifTags

# ANSI colors
RED = "\033[91m"
YELLOW = "\033[93m"
MAGENTA = "\033[95m"
CYAN = "\033[96m"
RESET = "\033[0m"
GREEN = "\033[92m"

print(fr'''
{RED}      
                                                                                                                  
 ,---.                                    ,--.                                               ,--.                 
/  .-' ,---. ,--.--. ,---. ,--,--,  ,---. `--' ,---.    ,--.--. ,---.  ,---.  ,---. ,--.--.,-'  '-. ,---. ,--.--. 
|  `-,| .-. ||  .--'| .-. :|      \(  .-' ,--.| .--'    |  .--'| .-. :| .-. || .-. ||  .--''-.  .-'| .-. :|  .--' 
|  .-'' '-' '|  |   \   --.|  ||  |.-'  `)|  |\ `--.    |  |   \   --.| '-' '' '-' '|  |     |  |  \   --.|  |    
`--'   `---' `--'    `----'`--''--'`----' `--' `---'    `--'    `----'|  |-'  `---' `--'     `--'   `----'`--'    
                                                                      `--'                                        
{YELLOW}|----------------------------------{MAGENTA}coded by Mishima and Nocturnis){YELLOW}---------------------------------------------|{RESET}
{CYAN}[+] • DOCX + CSV output •{GREEN}
''')

# -------------------- Config --------------------
SUPPORTED_EXT = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}
MAX_EMBED_PIXELS = 2500 * 2500  # avoid embedding absurdly huge images
EMBED_WIDTH_INCH = 1
CSV_ENCODING = 'utf-8-sig'
CSV_QUOTE = csv.QUOTE_ALL
EXIF_TAGS = {v: k for k, v in ExifTags.TAGS.items()}

# -------------------- Sorting Helper --------------------
def sort_paths_by_ext(paths):
    """Sort paths: .png first, then .jpg/.jpeg, then others. Stable alphabetical within groups."""
    ext_priority = {'.png': 0, '.jpg': 1, '.jpeg': 1}
    def sort_key(p):
        ext = Path(p).suffix.lower()
        priority = ext_priority.get(ext, 999)
        return (priority, str(p))
    return sorted(paths, key=sort_key)

# -------------------- Logging --------------------
logger = logging.getLogger("forensics_reporter")
handler = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s', '%Y-%m-%d %H:%M:%S')
handler.setFormatter(formatter)
logger.addHandler(handler)
logger.setLevel(logging.INFO)

# -------------------- Utilities --------------------
def sha256_data(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()

def format_ts(ts: float) -> str:
    return datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M:%S')

def safe_str(v) -> str:
    if v is None:
        return ''
    try:
        return str(v).replace('\x00', '').strip()
    except Exception:
        return ''

# -------------------- EXIF helpers --------------------
def get_exif_pillow_bytes(image_bytes: bytes) -> Dict[str, str]:
    try:
        with Image.open(io.BytesIO(image_bytes)) as img:
            exif = img.getexif()
            if not exif:
                return {}
            out = {}
            for k, v in exif.items():
                name = ExifTags.TAGS.get(k, k)
                out[name] = safe_str(v)
            return out
    except Exception as e:
        logger.debug(f"get_exif_pillow_bytes failed: {e}")
        return {}

def get_exif_pillow_path(path: str) -> Dict[str, str]:
    try:
        with Image.open(path) as img:
            exif = img.getexif()
            if not exif:
                return {}
            out = {}
            for k, v in exif.items():
                name = ExifTags.TAGS.get(k, k)
                out[name] = safe_str(v)
            return out
    except Exception as e:
        logger.debug(f"get_exif_pillow_path failed for {path}: {e}")
        return {}

# -------------------- Safe ZIP extraction --------------------
def safe_extract_zip(zf: zipfile.ZipFile, dest: str, members: Optional[list] = None) -> list:
    extracted = []
    dest = os.path.abspath(dest)
    names = members or zf.namelist()
    for name in names:
        if name.endswith('/') or name.startswith('__MACOSX'):
            continue
        normalized = os.path.normpath(name)
        if os.path.isabs(normalized) or normalized.startswith('..'):
            logger.warning(f"Skipping suspicious zip member (path traversal): {name}")
            continue
        ext = Path(normalized).suffix.lower()
        if ext not in SUPPORTED_EXT:
            continue
        target_path = os.path.join(dest, normalized)
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        with zf.open(name) as src, open(target_path, 'wb') as dst:
            dst.write(src.read())
        extracted.append(target_path)
    return extracted

# -------------------- Image compression --------------------
def fix_orientation(img: Image.Image) -> Image.Image:
    try:
        exif = img._getexif()
        if exif is None:
            return img
        orientation_key = next((k for k, v in ExifTags.TAGS.items() if v == 'Orientation'), None)
        if not orientation_key:
            return img
        orientation = exif.get(orientation_key)
        if orientation == 3:
            img = img.rotate(180, expand=True)
        elif orientation == 6:
            img = img.rotate(270, expand=True)
        elif orientation == 8:
            img = img.rotate(90, expand=True)
        return img
    except Exception:
        return img

def compress_image_bytes_for_docx(img_bytes: bytes, ext: str) -> bytes:
    try:
        with Image.open(io.BytesIO(img_bytes)) as img:
            img = fix_orientation(img)
            w, h = img.size
            pixels = w * h
            if pixels > MAX_EMBED_PIXELS:
                scale = (MAX_EMBED_PIXELS / pixels) ** 0.5
                new_w = max(100, int(w * scale))
                new_h = max(100, int(h * scale))
                img = img.resize((new_w, new_h), Image.LANCZOS)
            if img.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[-1])
                img = bg
            bio = io.BytesIO()
            img.save(bio, format='JPEG', quality=75, optimize=True)
            return bio.getvalue()
    except Exception as e:
        logger.debug(f"compress_image_bytes_for_docx failed: {e}")
        return img_bytes

# -------------------- DOCX helpers --------------------
def style_doc(doc: Document):
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Consola'
        font.size = Pt(9)
    except Exception:
        pass

def add_chain_of_custody(doc: Document):
    doc.add_heading('Chain of Custody', level=2)
    now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    p = doc.add_paragraph()
    p.add_run(f"Generated: {now}\n")
    p.add_run("Notes: Report generated by forensics_reporter.py")
    doc.add_paragraph('\n')

def add_image_table(doc: Document, image_bytes: bytes, metadata: Dict[str, str]):
    emb_bytes = compress_image_bytes_for_docx(image_bytes, metadata.get('ext', 'jpg'))
    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp:
        tmp.write(emb_bytes)
        tmp_path = tmp.name
    try:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        cell_img = table.cell(0, 0)
        paragraph = cell_img.paragraphs[0]
        run = paragraph.add_run()
        try:
            run.add_picture(tmp_path, width=Inches(EMBED_WIDTH_INCH))
        except Exception as e:
            cell_img.text = f"[Embed failed: {e}]"

        cell_meta = table.cell(0, 1)
        meta_lines = [f"SHA-256: {metadata.get('hash','')}"]
        if 'size' in metadata:
            meta_lines.append(f"Size: {metadata.get('size')} bytes")
        if 'modified' in metadata:
            meta_lines.append(f"Modified: {metadata.get('modified')}")
        if 'created' in metadata:
            meta_lines.append(f"Created*: {metadata.get('created')}")
        meta_lines.append(f"Source: {metadata.get('source','')}")

        if metadata.get('exif_summary'):
            meta_lines.append('\nEXIF:')
            meta_lines.extend(metadata['exif_summary'].split('\n'))

        cell_meta.text = '\n'.join(meta_lines)
        doc.add_paragraph()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

# -------------------- CSV helpers --------------------
def write_csv_header(csv_path: str):
    with open(csv_path, 'w', newline='', encoding=CSV_ENCODING) as f:
        writer = csv.writer(f, quoting=CSV_QUOTE)
        writer.writerow([
            'sha256', 'source', 'size_bytes', 'modified', 'created',
            'make', 'model', 'software', 'datetime_original'
        ])

def append_to_csv(csv_path: str, record: Dict[str, str]):
    with open(csv_path, 'a', newline='', encoding=CSV_ENCODING) as f:
        writer = csv.writer(f, quoting=CSV_QUOTE)
        writer.writerow([
            record.get('sha256',''),
            record.get('source',''),
            record.get('size_bytes',''),
            record.get('modified',''),
            record.get('created',''),
            record.get('make',''),
            record.get('model',''),
            record.get('software',''),
            record.get('datetime_original','')
        ])

# -------------------- Processing functions --------------------
def process_pdf(pdf_path: str, doc: Document, csv_path: str, seen_hashes: Set[str]):
    if not os.path.isfile(pdf_path):
        doc.add_paragraph(f"[!] PDF not found: {pdf_path}")
        return
    doc.add_heading('Source: Embedded Images from PDF', level=1)
    pdf_stat = os.stat(pdf_path)
    doc.add_paragraph(f"PDF Path: {os.path.abspath(pdf_path)}\nPDF Modified: {format_ts(pdf_stat.st_mtime)}\n")

    try:
        pdf_doc = fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"Failed to open PDF {pdf_path}: {e}")
        return

    count = 0
    try:
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            for img_idx, img in enumerate(page.get_images(full=True)):
                try:
                    xref = img[0]
                    base_img = pdf_doc.extract_image(xref)
                    img_bytes = base_img.get('image')
                    ext = base_img.get('ext') or 'jpg'
                    if not img_bytes:
                        continue
                    img_hash = sha256_data(img_bytes)
                    if img_hash in seen_hashes:
                        continue
                    seen_hashes.add(img_hash)
                    count += 1

                    exif = get_exif_pillow_bytes(img_bytes)
                    exif_keys = ['Make','Model','Software','DateTime','DateTimeOriginal']
                    exif_summary = '\n'.join(f"{k}: {exif.get(k,'')}" for k in exif_keys if exif.get(k))
                    if not exif_summary:
                        exif_summary = 'No EXIF data.'

                    metadata = {
                        'hash': img_hash,
                        'ext': ext,
                        'size': len(img_bytes),
                        'source': f"PDF Page {page_num+1}, Image {img_idx+1}",
                        'exif_summary': exif_summary
                    }

                    add_image_table(doc, img_bytes, metadata)

                    record = {
                        'sha256': img_hash,
                        'source': metadata['source'],
                        'size_bytes': len(img_bytes),
                        'modified': '',
                        'created': '',
                        'make': safe_str(exif.get('Make')),
                        'model': safe_str(exif.get('Model')),
                        'software': safe_str(exif.get('Software')),
                        'datetime_original': safe_str(exif.get('DateTimeOriginal') or exif.get('DateTime'))
                    }
                    append_to_csv(csv_path, record)
                except Exception as e:
                    logger.warning(f"Failed to extract an image from PDF page {page_num+1}: {e}")
    finally:
        pdf_doc.close()

    doc.add_paragraph(f"Extracted {count} unique images from PDF.\n")

# -------------------- Image processing --------------------
def process_single_image_file(img_path: str, doc: Document, csv_path: str, seen_hashes: Set[str], fs_warning: str, source_label: str = ''):
    try:
        img_hash = sha256_file(img_path)
        if img_hash in seen_hashes:
            return
        seen_hashes.add(img_hash)

        stat = os.stat(img_path)

        exif = get_exif_pillow_path(img_path)
        exif_time = exif.get('DateTimeOriginal') or exif.get('DateTime')
        time_warning = ''
        if exif_time:
            try:
                exif_dt = datetime.datetime.strptime(exif_time, '%Y:%m:%d %H:%M:%S')
                file_dt = datetime.datetime.fromtimestamp(stat.st_mtime)
                if abs((exif_dt - file_dt).total_seconds()) > 3600:
                    time_warning = 'EXIF/file timestamp mismatch (>1h) – possible spoofing.'
            except Exception:
                pass

        warning_full = (fs_warning + (' ' + time_warning if time_warning else '')).strip()

        with open(img_path, 'rb') as f:
            img_bytes = f.read()

        exif_summary = '\n'.join(f"{k}: {exif.get(k,'')}" for k in ['Make','Model','Software','DateTime','DateTimeOriginal'] if exif.get(k))
        if not exif_summary:
            exif_summary = 'No EXIF data.'

        metadata = {
            'hash': img_hash,
            'size': stat.st_size,
            'modified': format_ts(stat.st_mtime),
            'created': format_ts(stat.st_ctime),
            'source': source_label if source_label else img_path,
            'exif_summary': exif_summary,
            'warning': warning_full
        }

        add_image_table(doc, img_bytes, metadata)

        record = {
            'sha256': img_hash,
            'source': metadata['source'],
            'size_bytes': stat.st_size,
            'modified': format_ts(stat.st_mtime),
            'created': format_ts(stat.st_ctime),
            'make': safe_str(exif.get('Make')),
            'model': safe_str(exif.get('Model')),
            'software': safe_str(exif.get('Software')),
            'datetime_original': safe_str(exif_time)
        }
        append_to_csv(csv_path, record)

    except Exception as e:
        logger.warning(f"Failed to process image {img_path}: {e}")

def fast_file_hash(path: str, sample_size: int = 65536) -> str:
    """Fast hash using first + last sample_size bytes + file size."""
    stat = os.stat(path)
    size = stat.st_size
    h = hashlib.sha256()
    h.update(str(size).encode())
    with open(path, 'rb') as f:
        if size <= sample_size * 2:
            h.update(f.read())
        else:
            h.update(f.read(sample_size))
            f.seek(-sample_size, os.SEEK_END)
            h.update(f.read(sample_size))
    return h.hexdigest()

def process_images_source(input_path: str, doc: Document, csv_path: str, seen_hashes: Set[str]):
    fs_warning = "(*'Created' may reflect inode change time on Linux.)"
    image_paths = []

    if os.path.isdir(input_path):
        doc.add_heading('Source: Recovered Image Files (Directory)', level=1)
        image_paths = [str(p) for p in Path(input_path).rglob('*') if p.suffix.lower() in SUPPORTED_EXT]
        image_paths = sort_paths_by_ext(image_paths)
    elif zipfile.is_zipfile(input_path):
        doc.add_heading('Source: Recovered Image Files (ZIP)', level=1)
        with tempfile.TemporaryDirectory() as tmp:
            with zipfile.ZipFile(input_path, 'r') as zf:
                members = [m for m in zf.namelist() if Path(m).suffix.lower() in SUPPORTED_EXT and not m.endswith('/')]
                extracted = safe_extract_zip(zf, tmp, members=members)
                image_paths = sort_paths_by_ext(extracted)
    else:
        doc.add_paragraph(f"[!] Not valid dir/ZIP: {input_path}")
        return

    if not image_paths:
        doc.add_paragraph("[!] No supported images found.")
        return

    # === DEDUPLICATION PASS ===
    logger.info(f"Scanning {len(image_paths)} images for duplicates...")

    fast_hash_map = {}
    for path in image_paths:
        try:
            fhash = fast_file_hash(path)
            if fhash not in fast_hash_map:
                fast_hash_map[fhash] = []
            fast_hash_map[fhash].append(path)
        except Exception as e:
            logger.warning(f"Skipped {path} during fast hash: {e}")

    total_original = len(image_paths)
    kept_paths = []
    for group in fast_hash_map.values():
        if len(group) == 1:
            kept_paths.append(group[0])
        else:
            resolved = {}
            for p in group:
                try:
                    full_hash = sha256_file(p)
                    if full_hash not in resolved:
                        resolved[full_hash] = p
                except Exception as e:
                    logger.warning(f"Failed full hash on {p}: {e}")
            kept_paths.extend(resolved.values())

    duplicates_skipped = total_original - len(kept_paths)
    if duplicates_skipped > 0:
        logger.info(f"Skipped {duplicates_skipped} duplicate images.")

    # === SORT UNIQUE IMAGES BEFORE PROCESSING ===
    kept_paths = sort_paths_by_ext(kept_paths)

    # === PROCESS ONLY UNIQUE IMAGES ===
    for img_path in kept_paths:
        source_label = ''
        if zipfile.is_zipfile(input_path):
            arcname = os.path.relpath(img_path, os.path.commonpath([img_path, tmp]))
            source_label = f"{input_path} (member: {arcname})"
        process_single_image_file(img_path, doc, csv_path, seen_hashes, fs_warning, source_label)

# -------------------- Main --------------------
def main(pdf_path: Optional[str], files_path: Optional[str], output_docx: str):
    output_csv = str(Path(output_docx).with_suffix('.csv'))

    doc = Document()
    style_doc(doc)
    doc.add_heading('Forensics Evidence Report', 0)
    gen_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    doc.add_paragraph(f"Generated: {gen_time}\n")

    write_csv_header(output_csv)
    seen_hashes: Set[str] = set()

    if pdf_path:
        process_pdf(pdf_path, doc, output_csv, seen_hashes)
    else:
        doc.add_paragraph('No valid PDF provided.\n')

    if files_path:
        process_images_source(files_path, doc, output_csv, seen_hashes)
    else:
        doc.add_paragraph('No valid image folder/ZIP provided.\n')

    add_chain_of_custody(doc)

    try:
        doc.save(output_docx)
        logger.info(f"Report saved: {output_docx}")
        logger.info(f"Metadata CSV saved: {output_csv}")
    except Exception as e:
        logger.error(f"Failed to save report: {e}")

# -------------------- CLI --------------------
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Example: python forensics_reporter.py --pdf sample.pdf --images images_dir --out report.docx')
    parser.add_argument('--pdf', type=str, default=None, help='Path to PDF file (use None to skip)')
    parser.add_argument('--images', type=str, default=None, help='Path to directory or ZIP with images (use None to skip)')
    parser.add_argument('--out', type=str, required=True, help='Output DOCX path')

    args = parser.parse_args()

    if not args.pdf and not args.images:
        logger.error('Provide at least --pdf or --images')
        sys.exit(1)

    try:
        main(args.pdf, args.images, args.out)
    except KeyboardInterrupt:
        logger.info('Interrupted by user')
        sys.exit(2)
    except Exception as e:
        logger.exception(f"Unhandled exception: {e}")
        sys.exit(3)